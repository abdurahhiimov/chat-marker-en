#!/usr/bin/env python3
"""Builds the library of highlights in a Google Drive folder.

Does two things per run:
  1. picks up fresh exports from ~/Downloads into Inbox (backing up the previous one);
  2. rebuilds the structure — topics, pages, index, spreadsheet, dashboard.

Data model: a highlight has a topic (`tag`) and a color (`color`, plain hex, taken
from the topic). Old exports, where topics lived as hashtags in the note, still read.

From the terminal:
    python3 library.py              # normal run
    python3 library.py --no-ingest  # rebuild only
    python3 library.py --force      # rebuild even if nothing changed
"""

from __future__ import annotations

import json
import os
import re
import shutil
import sys
from collections import Counter, defaultdict
from datetime import datetime
from pathlib import Path
from typing import Any

# ---------------------------------------------------------------- structure

FOLDERS = {
    "inbox":    "00 Inbox",
    "topics":   "01 Topics",
    "chats":    "02 Pages",
    "untagged": "03 No topic",
    "drafts":   "04 Drafts",
    "docs":     "06 Documents",
    "archive":  "09 Archive",
}

# folders from older versions — cleaned up so they stop cluttering the place
LEGACY_FOLDERS = ["03 By meaning", "02 Chats", "05 Links"]

TAG_RE = re.compile(r"#([^\W\d_][\w-]{1,39})", re.UNICODE)
BACKUPS_KEEP = 30
NO_TAG = "no topic"
NO_TAG_COLOR = "#b0b3ba"
HEADER = "<!-- chatmarker:generated -->\n"
BUILD = "en-1"   # builder version: change it and the library is rebuilt by force


# ---------------------------------------------------------------- where the root is

def find_drive_root() -> Path:
    """Where the library lives: environment variable → root.txt → Drive autodetect.

    root.txt is written by the installer; the agent and MCP always went through it,
    but running library.py by hand didn't, and died with "Couldn't find Google Drive"
    even though the path had long been on disk. Now there is one source of truth
    for every entry point.
    """
    env = os.environ.get("CHATMARKER_ROOT", "").strip()
    if env:
        return Path(env).expanduser()

    stored = Path.home() / ".chatmarker" / "root.txt"
    try:
        if stored.is_file():
            first = stored.read_text(encoding="utf-8").strip().splitlines()
            if first and first[0].strip():
                return Path(first[0].strip()).expanduser()
    except OSError:
        pass

    home = Path.home()
    candidates: list[Path] = []

    cloud = home / "Library" / "CloudStorage"
    if cloud.is_dir():
        for d in sorted(cloud.glob("GoogleDrive-*")):
            for inner in ("My Drive",):
                if (d / inner).is_dir():
                    candidates.append(d / inner)

    for legacy in (home / "Google Drive" / "My Drive", home / "Google Drive"):
        if legacy.is_dir():
            candidates.append(legacy)

    if not candidates:
        raise SystemExit(
            "Couldn't find Google Drive on this Mac.\n"
            "Install Google Drive for desktop (google.com/drive/download), sign in,\n"
            "wait for the folder to appear and run the installer again.\n"
            "Or set the path by hand: CHATMARKER_ROOT=/path/to/folder python3 library.py"
        )
    return candidates[0] / "AI Highlights"


def ensure_tree(root: Path) -> dict[str, Path]:
    paths = {k: root / v for k, v in FOLDERS.items()}
    for p in paths.values():
        p.mkdir(parents=True, exist_ok=True)
    (paths["archive"] / "backups").mkdir(exist_ok=True)

    for old in LEGACY_FOLDERS:
        d = root / old
        if not d.is_dir() or d in paths.values():
            continue
        for f in list(d.glob("*.md")):
            try:
                if f.read_text(encoding="utf-8").startswith(HEADER.strip()):
                    f.unlink()
            except Exception:
                pass
        try:
            d.rmdir()          # goes only if empty — we don't touch your own stuff
        except OSError:
            pass
    return paths


# ---------------------------------------------------------------- taking data in

TOMB_KEEP = 5000     # how many tombstones the canonical file remembers


def _dicts(value) -> list[dict]:
    """A list of dicts out of anything — broken shapes are dropped silently."""
    if not isinstance(value, list):
        return []
    return [x for x in value if isinstance(x, dict)]


def _payload(data) -> dict:
    """A normalised export: any shape is brought to a single form."""
    if isinstance(data, list):
        data = {"highlights": data}
    if not isinstance(data, dict):
        return {"tags": [], "highlights": [], "deleted": [], "deletedTags": []}
    return {
        "tags": _dicts(data.get("tags")),
        "highlights": [h for h in _dicts(data.get("highlights")) if h.get("id")],
        "deleted": [str(x) for x in data.get("deleted") or [] if x],
        "deletedTags": [str(x).lower() for x in data.get("deletedTags") or [] if x],
    }


def merge_export(canon: dict, fresh: dict) -> dict:
    """Merging an export into the canonical file.

    The rules are simple: a record from the fresh export beats its own older version;
    records that aren't in the export stay (that's the buffer overflowing or another
    device, not a deletion); deletion happens only on an explicit tombstone.
    """
    by_id = {h["id"]: h for h in canon.get("highlights", [])}
    for h in fresh["highlights"]:
        by_id[h["id"]] = h

    tombs = set(canon.get("deleted") or []) | set(fresh["deleted"])
    for dead in tombs:
        by_id.pop(dead, None)

    by_name: dict[str, dict] = {}
    for t in canon.get("tags", []):
        name = str(t.get("name") or "").lower()
        if name:
            by_name[name] = t
    tag_tombs = set(canon.get("deletedTags") or []) | set(fresh["deletedTags"])
    for t in fresh["tags"]:
        name = str(t.get("name") or "").lower()
        if name:
            by_name[name] = t
            tag_tombs.discard(name)     # topic recreated — tombstone lifted
    for dead in tag_tombs:
        by_name.pop(dead, None)

    return {
        "version": 3,
        "tags": list(by_name.values()),
        "deletedTags": sorted(tag_tombs),
        "highlights": sorted(by_id.values(), key=lambda h: h.get("createdAt") or "", reverse=True),
        "deleted": sorted(tombs)[-TOMB_KEEP:],
    }


def write_atomic(path: Path, data: dict) -> None:
    tmp = path.with_suffix(".tmp")
    tmp.write_text(json.dumps(data, ensure_ascii=False, indent=1), encoding="utf-8")
    os.replace(tmp, path)


def ingest(root: Path, paths: dict[str, Path]) -> int:
    """Picks up exports from Downloads and MERGES them into the canonical file.

    An export used to overwrite the whole Inbox — two devices wiped each other,
    and whatever the buffer pushed out disappeared from the library. Now the Inbox
    is cumulative: no overwrite, just a merge by id.
    """
    downloads = Path.home() / "Downloads"
    try:
        if not downloads.is_dir():
            return 0
        # We look one level of subfolders deep too: "tidy up Downloads" is a very
        # natural urge, and it used to silently break the whole pipeline
        # (files piled up in ~/Downloads/Highlights while the library stayed empty).
        patterns = ("highlights*.json", "*/highlights*.json")
        files = sorted(
            [f for pat in patterns for f in downloads.glob(pat) if f.is_file()],
            key=lambda f: f.stat().st_mtime,
        )
    except OSError:      # no access to Downloads (TCC) — just rebuild without ingest
        return 0

    target = paths["inbox"] / "highlights.json"
    moved = 0
    for f in files:
        try:
            data = json.loads(f.read_text(encoding="utf-8"))
        except Exception:
            continue
        fresh = _payload(data)
        if not fresh["highlights"] and not fresh["deleted"] and not fresh["tags"]:
            continue

        canon = {}
        if target.exists():
            stamp = datetime.now().strftime("%Y%m%d-%H%M%S-%f")
            try:
                shutil.copy2(target, paths["archive"] / "backups" / f"highlights-{stamp}.json")
                canon = json.loads(target.read_text(encoding="utf-8"))
            except Exception:
                canon = {}
        if not isinstance(canon, dict):
            canon = {}

        write_atomic(target, merge_export(canon, fresh))
        try:
            f.unlink()
        except FileNotFoundError:
            pass         # another process got there first — fine, the merge is idempotent
        moved += 1

    backups = sorted((paths["archive"] / "backups").glob("highlights-*.json"),
                     key=lambda f: f.stat().st_mtime, reverse=True)
    for old in backups[BACKUPS_KEEP:]:
        old.unlink(missing_ok=True)
    return moved


def load_all(paths: dict[str, Path]) -> tuple[list[dict[str, Any]], dict[str, str], set[str]]:
    """Returns highlights, the topic palette (name -> color) and topics with a live doc."""
    items: list[dict[str, Any]] = []
    palette: dict[str, str] = {}
    docs: set[str] = set()
    decided: set[str] = set()     # the "keep a doc" flag comes from the freshest export
    seen: set[str] = set()

    files = sorted(paths["inbox"].glob("*.json"), key=lambda f: f.stat().st_mtime, reverse=True)
    for f in files:
        try:
            data = json.loads(f.read_text(encoding="utf-8"))
        except Exception:
            log(f"  ! could not read {f.name}, skipping")
            continue

        for t in _dicts(data.get("tags") if isinstance(data, dict) else None):
            name = str(t.get("name") or "").strip().lower()
            if not name:
                continue
            if t.get("color"):
                palette.setdefault(name, safe_color(t["color"]))
            if name not in decided:
                decided.add(name)
                if t.get("doc"):
                    docs.add(name)

        for h in _dicts(data if isinstance(data, list) else (data.get("highlights") if isinstance(data, dict) else None)):
            hid = h.get("id")
            if not hid or hid in seen:
                continue
            seen.add(hid)
            items.append(h)

    items.sort(key=lambda h: h.get("createdAt", ""), reverse=True)
    for h in items:
        for t in tags_of(h):
            palette.setdefault(t, color_of(h))
    return items, palette, docs


# ---------------------------------------------------------------- helpers

SPA_HOSTS = re.compile(r"(^|\.)(claude\.ai|chatgpt\.com|chat\.openai\.com|gemini\.google\.com|aistudio\.google\.com)$", re.I)


def fragment_url(h: dict) -> str:
    """A link that opens the source and scrolls to exactly this spot.

    The browser's own text fragments: page#:~:text=chunk
    We don't add context — start and end are enough, and gluing neighbouring
    blocks together in the index makes prefix/suffix unreliable.
    """
    raw = h.get("url") or ""
    url = raw.split("#")[0]
    if not url:
        return ""
    try:
        from urllib.parse import urlparse, quote
        parsed = urlparse(raw)
        if parsed.scheme == "file":
            return raw         # desktop highlight: file://…#page=N from the Viewer
        if parsed.scheme not in ("http", "https"):
            return ""          # javascript: and the like from untrusted exports
        if SPA_HOSTS.search(parsed.hostname or ""):
            return url
    except Exception:
        return ""

    ex = re.sub(r"\s+", " ", (h.get("anchor") or {}).get("exact") or h.get("text") or "").strip()
    if not ex:
        return url

    enc = lambda t: quote(t, safe="").replace("-", "%2D")
    if len(ex) > 70:
        head = re.sub(r"\s\S*$", "", ex[:40])
        tail = re.sub(r"^\S*\s", "", ex[-40:])
        directive = f"{enc(head)},{enc(tail)}"
    else:
        directive = enc(ex)
    return f"{url}#:~:text={directive}"


def fmt_of(h: dict[str, Any]) -> list[str]:
    """Formats of a highlight. New format is the fmt list, old one a single style."""
    f = h.get("fmt")
    if isinstance(f, list):
        return [x for x in f if x in ("b", "u", "s")]
    old = {"bold": ["b"], "under": ["u"], "strike": ["s"]}
    return old.get(h.get("style") or "", [])


def styled(h: dict[str, Any]) -> str:
    """How a highlight looks in text: formats are applied as markup."""
    t = (h.get("text") or "").strip()
    f = fmt_of(h)
    if "b" in f:
        t = f"**{t}**"
    if "u" in f:
        t = f"<u>{t}</u>"
    if "s" in f:
        t = f"~~{t}~~"
    return t


def tags_of(h: dict[str, Any]) -> list[str]:
    """Topics of a highlight. New format is the tag field, old one hashtags in the note."""
    raw = h.get("tag")
    if isinstance(raw, str) and raw.strip():
        return [raw.strip().lower()]
    slug = h.get("slug")
    if isinstance(slug, str) and slug.strip():
        return [slug.strip().lower()]
    many = h.get("tags")
    if isinstance(many, list) and many:
        return [str(t).strip().lower() for t in many if str(t).strip()]
    return sorted({m.group(1).lower() for m in TAG_RE.finditer(h.get("note") or "")})


HEX_RE = re.compile(r"#[0-9a-fA-F]{6}")

# Desktop captures (Hammerspoon) write the color as a word, not as hex.
# We map them onto the same shades as the browser palette — otherwise they went gray.
NAMED_COLORS = {
    "yellow": "#ffd54f",
    "green": "#81c784",
    "blue": "#64b5f6",
    "red": "#ef9a9a",
}


def safe_color(c) -> str:
    """A color out of an export is an untrusted string: only clean hex will do."""
    if isinstance(c, str):
        c = NAMED_COLORS.get(c.strip().lower(), c)
    return c if isinstance(c, str) and HEX_RE.fullmatch(c) else NO_TAG_COLOR


def color_of(h: dict[str, Any]) -> str:
    return safe_color(h.get("color") or "")


def slug(s: str, limit: int = 60) -> str:
    s = re.sub(r"[\\/:*?\"<>|\n\r\t]", " ", s or "").strip()
    s = re.sub(r"\s+", " ", s)
    return s[:limit].rstrip() or "untitled"


def entry_md(h: dict[str, Any], show_source: bool = True) -> str:
    tags = tags_of(h)
    out = ["> " + styled(h).replace("\n", "\n> "), ""]
    meta = ("`#" + "` `#".join(tags) + "`") if tags else f"`{NO_TAG}`"
    meta += f" · {(h.get('createdAt') or '')[:10]}"
    if show_source and h.get("title"):
        meta += f" · {h['title']}"
    out.append(meta)
    if h.get("note"):
        out.append(f"\n**Note:** {h['note']}")
    fu = fragment_url(h)
    if fu:
        out.append(f"\n[to this spot in the source]({fu})")
    out.append("\n---\n")
    return "\n".join(out)


def write(path: Path, content: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(content, encoding="utf-8")


def clean_generated(folder: Path) -> None:
    for f in folder.glob("*.md"):
        try:
            if f.read_text(encoding="utf-8").startswith(HEADER.strip()):
                f.unlink()
        except Exception:
            pass


# ---------------------------------------------------------------- build

def build_topics(items, paths) -> list[tuple[str, int]]:
    by_tag: dict[str, list[dict]] = defaultdict(list)
    for h in items:
        for t in tags_of(h):
            by_tag[t].append(h)

    clean_generated(paths["topics"])
    for tag, hs in by_tag.items():
        body = [HEADER, f"# #{tag}",
                f"_{len(hs)} {_plural(len(hs), 'highlight', 'highlights', 'highlights')}"
                f" · updated {datetime.now():%Y-%m-%d %H:%M}_", ""]
        by_src: dict[str, list[dict]] = defaultdict(list)
        for h in hs:
            by_src[h.get("title") or "untitled"].append(h)
        for src, chunk in by_src.items():
            body.append(f"## {src}\n")
            body += [entry_md(h, show_source=False) for h in chunk]
        write(paths["topics"] / f"{slug(tag)}.md", "\n".join(body))

    return sorted(((t, len(hs)) for t, hs in by_tag.items()), key=lambda x: -x[1])


def build_sources(items, paths) -> list[tuple[str, int, str]]:
    by_src: dict[str, list[dict]] = defaultdict(list)
    for h in items:
        by_src[h.get("title") or h.get("conv") or "untitled"].append(h)

    clean_generated(paths["chats"])
    rows = []
    for src, hs in by_src.items():
        last = max((h.get("createdAt") or "") for h in hs)
        url = next((h.get("url") for h in hs if h.get("url")), "")
        body = [HEADER, f"# {src}",
                f"_{len(hs)} {_plural(len(hs), 'highlight', 'highlights', 'highlights')}"
                f" · last one {last[:10]}_", ""]
        if url:
            body.append(f"[open the source]({url})\n")
        body += [entry_md(h, show_source=False) for h in hs]
        write(paths["chats"] / f"{last[:7]} {slug(src)}.md", "\n".join(body))
        rows.append((src, len(hs), last[:10]))
    return sorted(rows, key=lambda r: r[2], reverse=True)


def build_untagged(items, paths) -> int:
    hs = [h for h in items if not tags_of(h)]
    clean_generated(paths["untagged"])
    body = [HEADER, "# No topic",
            f"_{len(hs)} {_plural(len(hs), 'highlight', 'highlights', 'highlights')}"
            f" · updated {datetime.now():%Y-%m-%d %H:%M}_",
            "", "Sort them into topics in the browser side panel and they'll move out of here.", ""]
    body += [entry_md(h) for h in hs]
    write(paths["untagged"] / "To sort.md", "\n".join(body))
    return len(hs)


def build_index(items, paths, root, topics, sources, untagged, palette) -> None:
    total = len(items)
    with_note = sum(1 for h in items if h.get("note"))
    months = Counter((h.get("createdAt") or "")[:7] for h in items if h.get("createdAt"))
    hosts = Counter(h.get("host") or "" for h in items if h.get("host"))

    lines = [
        HEADER, "# Index of highlights",
        f"_updated {datetime.now():%Y-%m-%d %H:%M}_", "",
        f"In total **{total}**, with a note **{with_note}**, with no topic **{untagged}**, topics **{len(topics)}**.",
        "", "## Topics", "",
    ]
    if topics:
        lines += ["| topic | count | file |", "|---|---|---|"]
        for t, n in topics[:60]:
            lines.append(f"| #{t} | {n} | [{t}](<{FOLDERS['topics']}/{slug(t)}.md>) |")
    else:
        lines.append("_No topics yet. Select some text in the browser and make the first one right there._")

    if untagged:
        lines += ["", f"[Sort out the ones with no topic — {untagged}](<{FOLDERS['untagged']}/To sort.md>)"]

    lines += ["", "## Where from", "", "| source | count | last one |", "|---|---|---|"]
    for src, n, last in sources[:40]:
        lines.append(f"| {src} | {n} | {last} |")

    if hosts:
        lines += ["", "## Sites", ""]
        for host, n in hosts.most_common(15):
            lines.append(f"- {host} — {n}")

    if months:
        lines += ["", "## By month", ""]
        for m, n in sorted(months.items(), reverse=True)[:12]:
            lines.append(f"- {m} — {n}")

    lines += ["", "## Recent", ""]
    for h in items[:10]:
        lines.append(entry_md(h))

    write(root / "Index.md", "\n".join(lines))


CTRL_RE = re.compile(r"[\x00-\x08\x0b-\x0c\x0e-\x1f]")


def xls_safe(v):
    """openpyxl chokes on control characters — and they really do turn up
    in text copied off the web."""
    return CTRL_RE.sub(" ", v) if isinstance(v, str) else v


def build_xlsx(items, root, palette) -> bool:
    try:
        return _build_xlsx(items, root, palette)
    except Exception as e:          # file open in Excel, locked by Drive, and so on
        log(f"  ! the spreadsheet didn't build: {e}")
        return False


def _build_xlsx(items, root, palette) -> bool:
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Alignment, Font, PatternFill
        from openpyxl.utils import get_column_letter
    except ImportError:
        return False

    wb = Workbook()
    ws = wb.active
    ws.title = "Highlights"
    headers = ["Date", "Topic", "Highlight", "Note", "Source", "Site", "Link"]
    ws.append(headers)
    for c in range(1, len(headers) + 1):
        cell = ws.cell(row=1, column=c)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="3C4043")

    for h in items:
        tags = tags_of(h)
        ws.append([xls_safe(v) for v in (
            (h.get("createdAt") or "")[:10],
            ", ".join("#" + t for t in tags) or NO_TAG,
            styled(h),
            h.get("note", ""),
            h.get("title", ""),
            h.get("host", ""),
            h.get("url", ""),
        )])
        r = ws.max_row
        ws.cell(row=r, column=2).fill = PatternFill("solid", fgColor=color_of(h).lstrip("#"))
        for col in (3, 4):
            ws.cell(row=r, column=col).alignment = Alignment(wrap_text=True, vertical="top")
        fu = fragment_url(h)
        if fu:
            link = ws.cell(row=r, column=7)
            link.hyperlink = fu
            link.value = "open"
            link.font = Font(color="1155CC", underline="single")

    for col, width in zip("ABCDEFG", (12, 22, 70, 40, 30, 20, 10)):
        ws.column_dimensions[col].width = width
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = f"A1:{get_column_letter(len(headers))}{max(1, ws.max_row)}"
    wb.save(root / "Highlights.xlsx")
    return True


# ---------------------------------------------------------------- live documents

SESSION_GAP = 30 * 60          # a break over half an hour means a new session, as in the dashboard


def _sessions(hs: list[dict[str, Any]]) -> list[list[dict[str, Any]]]:
    """Highlights in order, split into sessions."""
    def ts(h):
        try:
            return datetime.fromisoformat((h.get("createdAt") or "").replace("Z", "+00:00")).timestamp()
        except Exception:
            return 0.0

    ordered = sorted(hs, key=lambda h: h.get("createdAt") or "")
    out: list[list[dict]] = []
    prev = None
    for h in ordered:
        t = ts(h)
        if prev is None or (t and prev and t - prev > SESSION_GAP):
            out.append([])
        out[-1].append(h)
        prev = t or prev
    return out


MONTHS = ["January", "February", "March", "April", "May", "June",
          "July", "August", "September", "October", "November", "December"]


def _plural(n: int, one: str, few: str, many: str) -> str:
    """English gets by with two forms: `few` and `many` are the same word here."""
    return one if n == 1 else few


def _session_title(h: dict[str, Any]) -> str:
    raw = h.get("createdAt") or ""
    try:
        d = datetime.fromisoformat(raw.replace("Z", "+00:00")).astimezone()
        return f"{MONTHS[d.month - 1]} {d.day}, {d:%H:%M}"
    except Exception:
        return raw[:16].replace("T", " ") or "no date"


def _left_bar(par, hex_color: str) -> None:
    """The colored bar on the left — the same trick as in the dashboard."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    pPr = par._p.get_or_add_pPr()
    bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), hex_color.lstrip("#").upper())
    bdr.append(left)
    pPr.append(bdr)


def _hyperlink(par, url: str, text: str) -> None:
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    r_id = par.part.relate_to(url, RT.HYPERLINK, is_external=True)
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color"); color.set(qn("w:val"), "1155CC"); rPr.append(color)
    und = OxmlElement("w:u"); und.set(qn("w:val"), "single"); rPr.append(und)
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "18"); rPr.append(sz)
    run.append(rPr)
    t = OxmlElement("w:t"); t.text = text
    run.append(t)
    link.append(run)
    par._p.append(link)


def topic_docx(tag: str, hs: list[dict[str, Any]], color: str, path: Path) -> None:
    """A live document for a topic: the same sequence as in the dashboard."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor, Inches

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Helvetica Neue"
    normal.font.size = Pt(11)

    title = doc.add_paragraph()
    run = title.add_run(f"#{tag}")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor.from_string(color.lstrip("#").upper())

    sub = doc.add_paragraph()
    word = _plural(len(hs), "highlight", "highlights", "highlights")
    srun = sub.add_run(f"{len(hs)} {word} · updated {datetime.now():%Y-%m-%d %H:%M}")
    srun.font.size = Pt(9)
    srun.font.color.rgb = RGBColor(0x8B, 0x8D, 0x96)
    note = sub.add_run("\nThe file rebuilds itself on every export — edits here get wiped.")
    note.font.size = Pt(9)
    note.italic = True
    note.font.color.rgb = RGBColor(0x8B, 0x8D, 0x96)

    n = 0
    for chunk in _sessions(hs):
        head = doc.add_paragraph()
        head.paragraph_format.space_before = Pt(18)
        hrun = head.add_run(_session_title(chunk[0]).upper())
        hrun.bold = True
        hrun.font.size = Pt(9)
        hrun.font.color.rgb = RGBColor(0x8B, 0x8D, 0x96)

        for h in chunk:
            n += 1
            par = doc.add_paragraph()
            par.paragraph_format.left_indent = Inches(0.22)
            par.paragraph_format.space_after = Pt(4)
            _left_bar(par, color)
            num = par.add_run(f"{n}. ")
            num.bold = True
            num.font.color.rgb = RGBColor.from_string(color.lstrip("#").upper())
            body = par.add_run((h.get("text") or "").strip())
            f = fmt_of(h)
            body.bold = "b" in f
            body.underline = "u" in f
            if "s" in f:
                body.font.strike = True

            if h.get("note"):
                np = doc.add_paragraph()
                np.paragraph_format.left_indent = Inches(0.32)
                np.paragraph_format.space_after = Pt(4)
                nr = np.add_run("✎ " + h["note"])
                nr.italic = True
                nr.font.size = Pt(10)
                nr.font.color.rgb = RGBColor(0xA1, 0x72, 0x2B)

            meta = doc.add_paragraph()
            meta.paragraph_format.left_indent = Inches(0.32)
            meta.paragraph_format.space_after = Pt(12)
            bits = " · ".join(x for x in [(h.get("title") or "").strip(), (h.get("createdAt") or "")[:10]] if x)
            mr = meta.add_run(bits + ("  " if bits else ""))
            mr.font.size = Pt(9)
            mr.font.color.rgb = RGBColor(0x8B, 0x8D, 0x96)
            fu = fragment_url(h)
            if fu:
                _hyperlink(meta, fu, "to this spot in the source")

    tail = doc.add_paragraph()
    tail.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = tail.add_run("built by Chat Marker")
    tr.font.size = Pt(8)
    tr.font.color.rgb = RGBColor(0xB0, 0xB3, 0xBA)

    doc.save(path)


def build_docs(items, paths, docs: set[str], palette) -> list[str]:
    """We only write a document for flagged topics. Other files are left alone."""
    folder = paths["docs"]
    folder.mkdir(parents=True, exist_ok=True)
    made: list[str] = []

    by_tag: dict[str, list[dict]] = defaultdict(list)
    for h in items:
        for t in tags_of(h):
            by_tag[t].append(h)

    try:
        import docx  # noqa: F401
        have_docx = True
    except ImportError:
        have_docx = False

    wanted: set[str] = set()
    for tag in sorted(docs):
        hs = by_tag.get(tag)
        if not hs:
            continue
        name = slug(tag)
        color = palette.get(tag, NO_TAG_COLOR)
        if have_docx:
            wanted.add(f"{name}.docx")
            try:
                topic_docx(tag, hs, color, folder / f"{name}.docx")
                made.append(tag)
            except Exception as e:                      # a document must not take the build down
                log(f"  ! the document for topic '{tag}' didn't build: {e}")
        else:
            wanted.add(f"{name}.md")
            body = [HEADER, f"# #{tag}",
                    f"_{len(hs)} {_plural(len(hs), 'highlight', 'highlights', 'highlights')}"
                    f" · updated {datetime.now():%Y-%m-%d %H:%M}_",
                    "", "_For .docx the python-docx library is missing:_",
                    "_`~/.chatmarker/venv/bin/pip install python-docx`_", ""]
            for chunk in _sessions(hs):
                body.append(f"## {_session_title(chunk[0])}\n")
                body += [entry_md(h, show_source=True) for h in chunk]
            write(folder / f"{name}.md", "\n".join(body))
            made.append(tag)

    # a topic you unchecked takes its own file with it
    known = {slug(t) for t in by_tag}
    for f in list(folder.glob("*.docx")) + list(folder.glob("*.md")):
        if f.name in wanted:
            continue
        if f.stem in known:
            f.unlink(missing_ok=True)

    readme = folder / "What's here.md"
    write(readme, f"""{HEADER}# Documents by topic

Live documents live here: one per topic you flagged in the browser with the
"keep a doc" button (highlights panel → topic list → document icon).

The file is rebuilt whole on every export, so edits inside it get wiped —
write your own thoughts in `{FOLDERS['drafts']}`.

Currently kept: {', '.join('#' + t for t in sorted(made)) if made else '— no topic is flagged'}.
""")
    return made


def build_dashboard(items, root, palette, docs=frozenset()) -> None:
    payload = [
        {
            "d": (h.get("createdAt") or "")[:10],
            "t": h.get("text", ""),
            "fm": fmt_of(h),
            "n": h.get("note", ""),
            "s": h.get("title", ""),
            # file://…#page=N (a desktop capture) has to arrive whole —
            # the fragment with the page number is the "to this spot" link
            "u": (h.get("url") or "") if str(h.get("url") or "").startswith("file:")
                 else (h.get("url") or "").split("#")[0],
            "g": tags_of(h),
            "c": color_of(h),
            "ts": h.get("createdAt") or "",
        }
        for h in items
    ]
    page = DASHBOARD_HTML
    def js(value) -> str:
        # </script> inside the data would close the script block — the dashboard's
        # only real hole. U+2028/2029 break strings in old engines — same deal.
        return (json.dumps(value, ensure_ascii=False)
                .replace("</", "<\\/")
                .replace("\u2028", "\\u2028").replace("\u2029", "\\u2029"))

    page = _icons(page)
    page = page.replace("__PALETTE__", js({k: safe_color(v) for k, v in palette.items()}))
    page = page.replace("__NOTAG__", js(NO_TAG_COLOR))
    page = page.replace("__DOCS__", js(sorted(docs)))
    page = page.replace("__DATA__", js(payload))
    write(root / "Dashboard.html", page)


DASHBOARD_HTML = """<!doctype html>
<html lang="en"><head><meta charset="utf-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>Highlights</title>
<style>
  :root { color-scheme: dark; --line:rgba(255,255,255,.1); --dim:#8b8d96; --fg:#eceef3;
          --glass:rgba(20,21,26,.66); }
  * { box-sizing:border-box; }
  body { margin:0; background:#0b0c0f; color:var(--fg);
         font:15px/1.6 -apple-system,BlinkMacSystemFont,"Segoe UI",Inter,sans-serif; }
  body::before { content:""; position:fixed; inset:-20%; z-index:-1; pointer-events:none;
    background:
      radial-gradient(50% 40% at 18% 10%, rgba(90,124,255,.20), transparent 60%),
      radial-gradient(45% 40% at 86% 20%, rgba(190,110,255,.15), transparent 60%),
      radial-gradient(50% 45% at 60% 94%, rgba(60,190,180,.12), transparent 60%); }
  svg.i { width:16px; height:16px; stroke:currentColor; fill:none; stroke-width:2;
          stroke-linecap:round; stroke-linejoin:round; vertical-align:-3px; }

  header { position:sticky; top:0; z-index:9; background:var(--glass);
           -webkit-backdrop-filter:blur(40px) saturate(190%); backdrop-filter:blur(40px) saturate(190%);
           border-bottom:1px solid var(--line); padding:14px 22px 12px; }
  .top { display:flex; align-items:center; gap:12px; margin-bottom:11px; }
  h1 { margin:0; font-size:18px; font-weight:650; letter-spacing:-.01em; }
  .sub { color:var(--dim); font-size:12.5px; }
  .grow { flex:1; }
  .searchbox { position:relative; margin-bottom:10px; }
  .searchbox svg { position:absolute; left:12px; top:11px; color:var(--dim); }
  input[type=search], input[type=text] { width:100%; background:rgba(0,0,0,.3);
      border:1px solid var(--line); color:var(--fg); border-radius:12px;
      padding:10px 12px 10px 36px; font:inherit; outline:none; }
  input[type=text] { padding-left:12px; }
  input:focus { border-color:rgba(120,150,255,.6); box-shadow:0 0 0 3px rgba(90,124,255,.16); }

  .chips { display:flex; gap:7px; flex-wrap:wrap; align-items:center; }
  .chip { border:1px solid rgba(255,255,255,.14); border-radius:999px; padding:5px 6px 5px 12px;
          font-size:12.5px; cursor:pointer; user-select:none; display:inline-flex;
          align-items:center; gap:8px; color:#dcdde3;
          background:linear-gradient(180deg,rgba(255,255,255,.11),rgba(255,255,255,.04));
          box-shadow:inset 0 1px 0 rgba(255,255,255,.16); transition:background .13s,transform .12s; }
  .chip:hover { background:linear-gradient(180deg,rgba(255,255,255,.2),rgba(255,255,255,.08)); }
  .chip:active { transform:scale(.97); }
  .chip.on { background:rgba(255,255,255,.2); color:#fff; border-color:rgba(255,255,255,.3); }
  .chip i.dot { width:9px; height:9px; border-radius:50%; }
  .chip .box { width:17px; height:17px; border-radius:6px; border:1.5px solid rgba(255,255,255,.32);
               display:inline-flex; align-items:center; justify-content:center; font-size:11px;
               color:transparent; }
  .chip .box.pin { background:#5a7cff; border-color:#5a7cff; color:#fff; }
  .chip.plain { padding:5px 13px; }
  .chip .doc { color:#9fb2ff; display:inline-flex; }
  .chip .doc svg.i { width:14px; height:14px; }

  .btn { border:1px solid var(--line); border-radius:12px; padding:8px 13px; font:inherit;
         font-size:12.5px; cursor:pointer; display:inline-flex; gap:7px; align-items:center;
         color:#dcdde3; background:linear-gradient(180deg,rgba(255,255,255,.11),rgba(255,255,255,.04));
         box-shadow:inset 0 1px 0 rgba(255,255,255,.16); }
  .btn:hover { background:linear-gradient(180deg,rgba(255,255,255,.2),rgba(255,255,255,.08)); color:#fff; }
  .btn.primary { background:rgba(90,124,255,.92); border-color:transparent; color:#fff; }
  .btn.primary:hover { background:rgba(104,138,255,1); }

  .pop { position:absolute; z-index:20; margin-top:8px; width:340px; max-width:92vw;
         background:var(--glass); border:1px solid rgba(255,255,255,.16); border-radius:18px;
         padding:9px; box-shadow:0 24px 60px rgba(0,0,0,.5);
         -webkit-backdrop-filter:blur(40px) saturate(190%); backdrop-filter:blur(40px) saturate(190%); }
  .rows { max-height:270px; overflow-y:auto; display:flex; flex-direction:column; gap:2px;
          margin-top:8px; padding:3px; border-radius:13px; background:rgba(0,0,0,.24);
          border:1px solid rgba(255,255,255,.07); }
  .rows::-webkit-scrollbar { width:8px; }
  .rows::-webkit-scrollbar-thumb { background:rgba(255,255,255,.16); border-radius:8px; }
  .row { display:flex; align-items:center; gap:10px; height:36px; padding:0 10px; border:0;
         border-radius:9px; background:none; color:#e9e9ee; font:inherit; font-size:13px;
         text-align:left; cursor:pointer; width:100%; }
  .row:hover { background:rgba(255,255,255,.1); }
  .row .nm { flex:1; overflow:hidden; text-overflow:ellipsis; white-space:nowrap; }
  .row .ct { color:#8b8d96; font-size:11.5px; }
  .row .box { width:18px; height:18px; border-radius:6px; border:1.5px solid rgba(255,255,255,.3);
              display:inline-flex; align-items:center; justify-content:center; font-size:11px;
              color:transparent; flex:none; }
  .row .box.pin { background:#5a7cff; border-color:#5a7cff; color:#fff; }

  main { padding:18px 22px 90px; max-width:940px; margin:0 auto; }
  .session { display:flex; align-items:center; gap:10px; margin:26px 0 12px; color:var(--dim);
             font-size:11.5px; text-transform:uppercase; letter-spacing:.06em; }
  .session::after { content:""; flex:1; height:1px; background:var(--line); }
  .session:first-child { margin-top:4px; }

  .item { display:flex; gap:14px; margin-bottom:10px; }
  .num { flex:none; width:28px; height:28px; border-radius:50%; display:flex; align-items:center;
         justify-content:center; font-size:11.5px; font-weight:700; color:#0d0e11; margin-top:3px; }
  .body { flex:1; background:rgba(255,255,255,.045); border:1px solid var(--line);
          border-left:3px solid #555; border-radius:16px; padding:13px 16px;
          -webkit-backdrop-filter:blur(18px); backdrop-filter:blur(18px);
          box-shadow:0 8px 26px rgba(0,0,0,.24); }
  .q { white-space:pre-wrap; }
  .f-b { font-weight:700; } .f-u { text-decoration:underline; text-underline-offset:3px; }
  .f-s { text-decoration:line-through; } .f-u.f-s { text-decoration:underline line-through; }
  .note { margin-top:8px; color:#ffd9a0; font-size:13px; white-space:pre-wrap; }
  .meta { margin-top:10px; font-size:11.5px; color:var(--dim); display:flex; gap:12px;
          flex-wrap:wrap; align-items:center; }
  .meta a { color:#8fa8ff; text-decoration:none; display:inline-flex; gap:5px; align-items:center; }
  .meta button { background:none; border:0; color:var(--dim); font:inherit; cursor:pointer;
                 padding:0; display:inline-flex; gap:5px; align-items:center; }
  .meta button:hover, .meta a:hover { color:#fff; }
  .tag { cursor:pointer; }
  .empty { color:#63656d; padding:60px 0; text-align:center; }
  .btn.more { display:block; margin:22px auto 0; padding:10px 22px; }
  .toast { position:fixed; left:50%; bottom:26px; transform:translateX(-50%); background:var(--glass);
           border:1px solid rgba(255,255,255,.16); border-radius:999px; padding:9px 18px;
           -webkit-backdrop-filter:blur(30px); backdrop-filter:blur(30px); z-index:30; }
</style></head>
<body>
<header>
  <div class="top">
    <h1>Highlights</h1>
    <div class="sub" id="sub"></div>
    <span class="grow"></span>
    <button class="btn" id="copyAll">__ICON_COPY__ Copy</button>
    <button class="btn" id="dl">__ICON_DL__ Download as is</button>
    <button class="btn primary" id="ask">__ICON_AI__ Ask Claude to sort it out</button>
  </div>
  <div class="searchbox">__ICON_SEARCH__<input type="search" id="q" placeholder="Search text, notes, sources"></div>
  <div class="chips" id="tags"></div>
</header>
<main id="view"></main>
<script>
const DATA = __DATA__;
const PALETTE = __PALETTE__;
const NOTAG = __NOTAG__;
const DOCS = new Set(__DOCS__);   // topics that have a live document
const GAP = 30 * 60 * 1000;      // a break over half an hour means a new session

/* Filter: clicking the name switches to that topic, the checkbox pins it.
   Pinned ones add up, the current one is added on top of them. */
const state = { q:"", current:null, pinned:new Set(), menu:false };

const counts = {};
DATA.forEach(h => (h.g.length ? h.g : ["no topic"]).forEach(t => counts[t] = (counts[t]||0)+1));
const lastSeen = {};
DATA.forEach(h => (h.g.length ? h.g : ["no topic"]).forEach(t => {
  if (!lastSeen[t] || h.ts > lastSeen[t]) lastSeen[t] = h.ts; }));
const byRecent = Object.keys(counts).sort((a,b) => (lastSeen[b]||"").localeCompare(lastSeen[a]||""));
const colorOf = t => PALETTE[t] || NOTAG;

/* The "to this spot" link: a browser text fragment. Computed on the fly —
   storing percent-encoded text in the file was many times heavier than the text. */
const SPA_HOSTS = /(^|\\.)(claude\\.ai|chatgpt\\.com|chat\\.openai\\.com|gemini\\.google\\.com|aistudio\\.google\\.com)$/i;
function fragUrl(h) {
  const base = h.u || "";
  if (!base) return "";
  try {
    const u = new URL(base);
    if (u.protocol === "file:") return base;   // desktop highlight: file://…#page=N
    if (u.protocol !== "http:" && u.protocol !== "https:") return "";
    if (SPA_HOSTS.test(u.hostname)) return base;
  } catch { return ""; }
  const ex = String(h.t || "").replace(/\\s+/g, " ").trim();
  if (!ex) return base;
  const enc = t => encodeURIComponent(t).replace(/-/g, "%2D");
  let dir;
  if (ex.length > 70) {
    const head = ex.slice(0, 40).replace(/\\s\\S*$/, "");
    const tail = ex.slice(-40).replace(/^\\S*\\s/, "");
    dir = enc(head) + "," + enc(tail);
  } else dir = enc(ex);
  return base + "#:~:text=" + dir;
}
const esc = s => String(s).replace(/[&<>"]/g, c => ({"&":"&amp;","<":"&lt;",">":"&gt;",'"':"&quot;"}[c]));

function active() {
  const set = new Set(state.pinned);
  if (state.current) set.add(state.current);
  return set;
}

function match(h) {
  const on = active();
  if (on.size) {
    const own = h.g.length ? h.g : ["no topic"];
    if (!own.some(t => on.has(t))) return false;
  }
  if (state.q) {
    const blob = (h.t + " " + h.n + " " + h.s).toLowerCase();
    return state.q.toLowerCase().split(/\\s+/).every(w => blob.includes(w));
  }
  return true;
}
const rows = () => DATA.filter(match).sort((a,b) => (a.ts < b.ts ? -1 : a.ts > b.ts ? 1 : 0));

function chipHtml(t) {
  const on = active().has(t), pin = state.pinned.has(t);
  return `<span class="chip ${on?"on":""}" data-t="${esc(t)}">
      <i class="dot" style="background:${colorOf(t)}"></i>${esc(t)}
      ${DOCS.has(t) ? `<span class="doc" title="A document is kept for this topic">__ICON_DOC__</span>` : ""}
      <span class="box ${pin?"pin":""}" data-pin="${esc(t)}" title="Pin">✓</span></span>`;
}

function chips() {
  const shown = byRecent.slice(0, 5);
  document.getElementById("tags").innerHTML =
    `<span class="chip plain ${active().size?"":"on"}" data-all="1">All</span>`
    + shown.map(chipHtml).join("")
    + (byRecent.length > 5 ? `<span class="chip plain" data-menu="1" style="font-size:16px;letter-spacing:1px">⋯</span>` : "")
    + (state.menu ? menuHtml() : "");
}

function menuHtml(filter) {
  const q = (filter || "").toLowerCase();
  const list = byRecent.filter(t => t.includes(q));
  return `<div class="pop" id="menu">
      <input type="text" id="menuq" placeholder="Find a topic…" value="${esc(filter||"")}">
      <div class="rows">${list.length ? list.map(t => `
        <button class="row" data-t="${esc(t)}">
          <i class="dot" style="background:${colorOf(t)};width:9px;height:9px;border-radius:50%"></i>
          <span class="nm">${esc(t)}</span><span class="ct">${counts[t]}</span>
          <span class="box ${state.pinned.has(t)?"pin":""}" data-pin="${esc(t)}" title="Pin">✓</span>
        </button>`).join("") : '<div class="empty" style="padding:24px 0">nothing found</div>'}</div>
    </div>`;
}

function fmtSession(ts) {
  const d = new Date(ts);
  const day = d.toLocaleDateString("en-US", { day:"numeric", month:"long" });
  const time = d.toLocaleTimeString("en-US", { hour:"2-digit", minute:"2-digit", hour12:false });
  return `${day}, ${time}`;
}

const PAGE_SIZE = 400;                    // cards at a time: long feeds are rendered in chunks
let shownCount = PAGE_SIZE;

function itemHtml(h, i) {
  return `<div class="item">
      <div class="num" style="background:${h.c}">${i+1}</div>
      <div class="body" style="border-left-color:${h.c}">
        <div class="q ${h.fm.map(x=>"f-"+x).join(" ")}">${esc(h.t)}</div>
        ${h.n ? `<div class="note">✎ ${esc(h.n)}</div>` : ""}
        <div class="meta">
          ${h.g.map(x=>`<span class="tag" data-t="${esc(x)}" style="color:${colorOf(x)}">#${esc(x)}</span>`).join(" ")}
          <span>${esc(h.s)}</span>
          <span>${h.d}</span>
          ${fragUrl(h) ? `<a href="${esc(fragUrl(h))}" target="_blank" rel="noopener">__ICON_EXT__ to the spot</a>` : ""}
          <button data-copy="${i}">__ICON_COPY__ copy</button>
        </div>
      </div></div>`;
}

function render(keepShown) {
  if (!keepShown) shownCount = PAGE_SIZE;
  const r = rows();
  const on = [...active()];
  document.getElementById("sub").textContent =
    `${r.length} of ${DATA.length}` + (on.length ? ` · ${on.join(", ")}` : ` · ${byRecent.length} topics`);

  if (!r.length) { document.getElementById("view").innerHTML = '<div class="empty">Nothing found</div>'; return; }

  const slice = r.slice(0, shownCount);
  let html = "", prev = null;
  slice.forEach((h, i) => {
    const t = Date.parse(h.ts) || 0;
    if (prev === null || t - prev > GAP) html += `<div class="session">${fmtSession(h.ts)}</div>`;
    prev = t;
    html += itemHtml(h, i);
  });
  if (r.length > slice.length) {
    html += `<button class="btn more" id="more">Show ${Math.min(PAGE_SIZE, r.length - slice.length)} more of ${r.length - slice.length}</button>`;
  }
  document.getElementById("view").innerHTML = html;
  document.getElementById("view")._rows = r;
}

function toast(msg) {
  const el = document.createElement("div");
  el.className = "toast"; el.textContent = msg;
  document.body.appendChild(el);
  setTimeout(() => el.remove(), 1600);
}

function asText(list) {
  const on = [...active()];
  const head = on.length ? `# ${on.join(", ")}\\n\\n` : "# All highlights\\n\\n";
  return head + list.map((h,i) =>
    `${i+1}. ${h.t}` + (h.n ? `\\n   note: ${h.n}` : "") + (fragUrl(h) ? `\\n   ${fragUrl(h)}` : "")
  ).join("\\n\\n") + "\\n";
}

document.addEventListener("click", e => {
  const pin = e.target.closest("[data-pin]");
  if (pin) {
    e.stopPropagation();
    const t = pin.dataset.pin;
    state.pinned.has(t) ? state.pinned.delete(t) : state.pinned.add(t);
    if (state.current === t) state.current = null;
    chips(); return render();
  }
  if (e.target.closest("[data-all]")) { state.current = null; state.pinned.clear(); state.menu = false; chips(); return render(); }
  if (e.target.closest("[data-menu]")) { state.menu = !state.menu; chips(); return; }

  const t = e.target.closest("[data-t]");
  if (t) {
    state.current = state.current === t.dataset.t ? null : t.dataset.t;
    state.menu = false; chips(); return render();
  }
  if (!e.target.closest("#menu")) { if (state.menu) { state.menu = false; chips(); } }

  if (e.target.closest("#more")) { shownCount += PAGE_SIZE; return render(true); }
  if (e.target.closest("#ask")) return askClaude();
  if (e.target.closest("#dl")) return downloadPlain();
  if (e.target.closest("#copyAll")) {
    navigator.clipboard.writeText(asText(rows())); return toast("Topic copied");
  }
  const cp = e.target.closest("[data-copy]");
  if (cp) {
    const h = document.getElementById("view")._rows[+cp.dataset.copy];
    navigator.clipboard.writeText("> " + h.t + (h.n ? "\\n\\n" + h.n : "") + (fragUrl(h) ? "\\n\\n" + fragUrl(h) : ""));
    toast("Copied");
  }
});

document.addEventListener("input", e => {
  if (e.target.id === "menuq") {
    const box = document.getElementById("menu");
    const v = e.target.value;
    box.outerHTML = menuHtml(v);
    const inp = document.getElementById("menuq");
    inp.focus(); inp.setSelectionRange(v.length, v.length);
    return;
  }
  if (e.target.id === "q") { state.q = e.target.value; qDebounce(); }
});

function topicName() {
  const on = [...active()];
  return on.length ? on.join(", ") : (state.q || "all topics");
}

function fileName(ext) {
  const d = new Date().toISOString().slice(0,10);
  const t = topicName().replace(/[\\\\/:*?"<>|]/g, " ").trim().slice(0, 60) || "highlights";
  return `${t} — ${d}.${ext}`;
}

/* Download as is: a markdown digest of whatever is in the filter right now.
   No processing — just the sequence, the way it was marked. */
function asDoc(list) {
  const head = `# ${topicName()}\\n\\n_${list.length} highlight${list.length === 1 ? "" : "s"} · built ${new Date().toLocaleDateString("en-US")}_\\n\\n`;
  let out = head, prev = null;
  list.forEach((h, i) => {
    const t = Date.parse(h.ts) || 0;
    if (prev === null || t - prev > GAP) out += `\\n## ${fmtSession(h.ts)}\\n\\n`;
    prev = t;
    out += `**${i+1}.** ${h.t}\\n\\n`;
    if (h.n) out += `> ✎ ${h.n}\\n\\n`;
    const meta = [h.s, h.d].filter(Boolean).join(" · ");
    if (meta) out += `<sub>${meta}</sub>`;
    const fu = fragUrl(h);
    if (fu) out += `${meta ? " · " : ""}[to the spot](${fu})`;
    out += "\\n\\n";
  });
  return out;
}

function downloadPlain() {
  const r = rows();
  if (!r.length) return toast("Nothing to download");
  const blob = new Blob([asDoc(r)], { type: "text/markdown;charset=utf-8" });
  const a = document.createElement("a");
  a.href = URL.createObjectURL(blob);
  a.download = fileName("md");
  document.body.appendChild(a);
  a.click();
  a.remove();
  setTimeout(() => URL.revokeObjectURL(a.href), 4000);
  toast("File saved to Downloads");
}

/* Ask Claude to sort it out: we open Claude with the task of building a document.
   It pulls the data itself through MCP highlights — here we only pass the first
   highlights, so the conversation starts with context even if MCP isn't connected. */
function askClaude() {
  const r = rows();
  if (!r.length) return toast("Nothing to sort out");
  const topic = topicName();
  let q = `Put together a digest on the topic "${topic}" from my highlights.\\n\\n`
        + `1. Take the full sequence through MCP highlights: get_topic_map(topic: "${topic}"), and search_highlights if that isn't enough.\\n`
        + `2. Sort it out: where I was heading, what repeats, where the contradictions are, what's missing. Keep the quotes word for word, keep the "to this spot" links.\\n`
        + `3. Put the finished text into a Google Doc "${topic} — digest" in the AI Highlights folder through the Google Drive connector, then give me the link.\\n\\n`
        + `Here are those same highlights in the order I marked them (${r.length}):\\n`
        + r.slice(0,60).map((h,i) => `${i+1}. ${h.t}` + (h.n ? ` — note: ${h.n}` : "")).join("\\n");
  if (q.length > 13000) q = q.slice(0, 13000);
  location.href = "claude://claude.ai/new?q=" + encodeURIComponent(q);
}

let qTimer = null;
function qDebounce() { clearTimeout(qTimer); qTimer = setTimeout(() => render(), 160); }

chips(); render();
</script></body></html>"""


# Icons drawn by hand in the Tabler manner: 24 grid, stroke 2, rounded caps.
ICONS = {
    "__ICON_SEARCH__": '<circle cx="10.5" cy="10.5" r="6.5"/><path d="M15.5 15.5 21 21"/>',
    "__ICON_LIST__":   '<path d="M8 6h13M8 12h13M8 18h13M3.5 6h.01M3.5 12h.01M3.5 18h.01"/>',
    "__ICON_MAP__":    '<circle cx="5" cy="6" r="2.5"/><circle cx="5" cy="18" r="2.5"/><circle cx="19" cy="12" r="2.5"/><path d="M7.3 7.2 16.7 11M7.3 16.8 16.7 13"/>',
    "__ICON_AI__":     '<path d="M12 3l1.8 4.6L18.5 9l-4.7 1.4L12 15l-1.8-4.6L5.5 9l4.7-1.4z"/><path d="M18 16l.9 2.2L21 19l-2.1.8L18 22l-.9-2.2L15 19l2.1-.8z"/>',
    "__ICON_EXT__":    '<path d="M11 5H6a2 2 0 0 0-2 2v11a2 2 0 0 0 2 2h11a2 2 0 0 0 2-2v-5"/><path d="M14 4h6v6M20 4 11 13"/>',
    "__ICON_COPY__":   '<rect x="8" y="8" width="12" height="12" rx="2"/><path d="M16 8V6a2 2 0 0 0-2-2H6a2 2 0 0 0-2 2v8a2 2 0 0 0 2 2h2"/>',
    "__ICON_DL__":     '<path d="M12 4v12M7 11l5 5 5-5M4 20h16"/>',
    "__ICON_DOC__":    '<path d="M13 3H7a2 2 0 0 0-2 2v14a2 2 0 0 0 2 2h10a2 2 0 0 0 2-2V9z"/><path d="M13 3v6h6M9 13h6M9 17h4"/>',
}


def _icons(page: str) -> str:
    """We insert icons with single quotes: they land inside JS strings that use
    double quotes, and they must not break anything there."""
    for key, body in ICONS.items():
        svg = f"<svg class='i' viewBox='0 0 24 24'>{body}</svg>".replace('"', "'")
        page = page.replace(key, svg)
    return page


def build_readme(root: Path) -> None:
    p = root / "How this all works.md"
    write(p, f"""{HEADER}# How this all works

Chat Marker fills this folder. Edit only `{FOLDERS['drafts']}` by hand — everything else
is rebuilt on every export, and edits there get wiped.

- `{FOLDERS['inbox']}` — raw data, the only source. Don't delete.
- `{FOLDERS['topics']}` — collections by topic. A topic is created in the browser as you select text.
- `{FOLDERS['chats']}` — the same, but laid out by the pages and chats it came from.
- `{FOLDERS['untagged']}` — whatever was saved with no topic. Sort it out when you get to it.
- `{FOLDERS['docs']}` — live documents for flagged topics, they rebuild themselves.
- `{FOLDERS['drafts']}` — your territory.
- `{FOLDERS['archive']}` — past versions of the export, the last {BACKUPS_KEEP}.
- `Index.md` — where to start.
- `Highlights.xlsx` — open it in Google Sheets, filters on every column.
- `Dashboard.html` — open it with a double click, works without the internet.
""")


# ---------------------------------------------------------------- running

def needs_rebuild(root: Path, paths: dict[str, Path]) -> bool:
    index = root / "Index.md"
    if not index.exists():
        return True
    # the builder was updated — rebuild even if the highlights are the same.
    # Without this a new dashboard wouldn't show up until the next export.
    stamp = paths["archive"] / ".build"
    try:
        if stamp.read_text(encoding="utf-8").strip() != BUILD:
            return True
    except OSError:
        return True
    try:
        built = index.stat().st_mtime
        return any(f.stat().st_mtime > built for f in paths["inbox"].glob("*.json"))
    except OSError:
        return True


class _Lock:
    """A lock between the launchd agent and MCP: building at the same time is a race.

    The lock file lives in the home folder, not on Google Drive — network mounts
    handle flock badly.
    """

    def __init__(self):
        self.fh = None

    def __enter__(self):
        try:
            import fcntl
            lock_dir = Path.home() / ".chatmarker"
            lock_dir.mkdir(parents=True, exist_ok=True)
            self.fh = open(lock_dir / ".library.lock", "w")
            fcntl.flock(self.fh, fcntl.LOCK_EX)
        except Exception:
            self.fh = None     # no fcntl or no rights — work without the lock, as before
        return self

    def __exit__(self, *a):
        if self.fh:
            try:
                import fcntl
                fcntl.flock(self.fh, fcntl.LOCK_UN)
                self.fh.close()
            except Exception:
                pass


def log(msg: str) -> None:
    """Service messages go to stderr: the MCP process's stdout is taken by the protocol."""
    print(msg, file=sys.stderr)


def run(do_ingest: bool = True, force: bool = False, quiet: bool = False) -> dict:
    with _Lock():
        return _run_locked(do_ingest, force, quiet)


def _run_locked(do_ingest: bool = True, force: bool = False, quiet: bool = False) -> dict:
    """One run: taking exports in plus a rebuild. Also used from the MCP server."""
    root = find_drive_root()
    paths = ensure_tree(root)

    moved = ingest(root, paths) if do_ingest else 0
    result = {"root": str(root), "moved": moved, "rebuilt": False, "skipped": False,
              "items": 0, "topics": 0, "sources": 0, "untagged": 0, "docs": [], "xlsx": False}

    if not force and not moved and not needs_rebuild(root, paths):
        result["skipped"] = True
        return result

    items, palette, docs = load_all(paths)
    result["items"] = len(items)
    if not items:
        build_readme(root)
        (paths["archive"] / ".build").write_text(BUILD, encoding="utf-8")
        return result

    topics = build_topics(items, paths)
    sources = build_sources(items, paths)
    untagged = build_untagged(items, paths)
    build_index(items, paths, root, topics, sources, untagged, palette)
    build_dashboard(items, root, palette, docs)
    made = build_docs(items, paths, docs, palette)
    build_readme(root)

    (paths["archive"] / ".build").write_text(BUILD, encoding="utf-8")
    result.update({
        "topics": len(topics), "sources": len(sources), "untagged": untagged,
        "docs": made, "xlsx": build_xlsx(items, root, palette), "rebuilt": True,
    })
    return result


def main() -> int:
    quiet = "--quiet" in sys.argv
    r = run(do_ingest="--no-ingest" not in sys.argv, force="--force" in sys.argv, quiet=quiet)
    if quiet:
        return 0

    if r["moved"]:
        print(f"  exports taken in: {r['moved']}")
    if r["skipped"]:
        print("  everything is built already, nothing changed")
    elif not r["rebuilt"]:
        print("  no highlights yet, the folder is ready")
    else:
        print(f"  built: {r['items']} highlights, {r['topics']} topics, {r['sources']} sources"
              + (f", no topic {r['untagged']}" if r["untagged"] else "")
              + (f", documents {len(r['docs'])}" if r["docs"] else "")
              + ("" if r["xlsx"] else " (no xlsx — openpyxl missing)"))
    print(f"  folder: {r['root']}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
